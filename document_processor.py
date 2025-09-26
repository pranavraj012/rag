# document_processor.py
import os
import json
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading TXT {file_path}: {e}")
            return ""
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Process single document and return chunks"""
        file_extension = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            text = self.extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file type: {file_extension}")
            return []
        
        if not text.strip():
            return []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangchainDocument objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": str(file_path),
                    "chunk_id": i,
                    "file_name": Path(file_path).name,
                    "file_type": file_extension
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_folder(self, folder_path: str) -> List[LangchainDocument]:
        """Process all documents in a folder"""
        all_documents = []
        supported_extensions = {'.pdf', '.docx', '.txt'}
        
        folder = Path(folder_path)
        if not folder.exists():
            folder.mkdir(parents=True, exist_ok=True)
            return []
        
        for file_path in folder.rglob("*"):
            if file_path.suffix.lower() in supported_extensions and file_path.is_file():
                print(f"Processing: {file_path.name}")
                documents = self.process_document(str(file_path))
                all_documents.extend(documents)
        
        return all_documents
